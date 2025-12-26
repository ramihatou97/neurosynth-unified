"""
Synthesis Export Module - Publication-Ready Output Generation

Unified export interface that routes to the optimal generator for each format:
- PDF: MedicalPDFGenerator (ReportLab) - native embedding, lossless quality
- HTML: Self-contained with base64 images
- DOCX: python-docx with embedded images
- Markdown: Figure links resolved

Usage:
    from synthesis_export import SynthesisExporter, ExportConfig
    
    exporter = SynthesisExporter(image_base_path=Path("data/images"))
    exporter.to_pdf(synthesis_result, Path("output.pdf"))
"""

import re
import base64
import json
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO

logger = logging.getLogger(__name__)

# =============================================================================
# IMPORT SPECIALIZED GENERATORS
# =============================================================================

# Import MedicalPDFGenerator (ReportLab-based)
_MedicalPDFGenerator = None
_PDFConfig = None

def _get_pdf_generator():
    """Lazy import of MedicalPDFGenerator to avoid import errors if not needed."""
    global _MedicalPDFGenerator, _PDFConfig
    
    if _MedicalPDFGenerator is not None:
        return _MedicalPDFGenerator, _PDFConfig
    
    # Try various import paths
    import_paths = [
        ("src.synthesis.medical_pdf_generator", "MedicalPDFGenerator", "PDFConfig"),
        ("src.medical_pdf_generator", "MedicalPDFGenerator", "PDFConfig"),
        ("medical_pdf_generator", "MedicalPDFGenerator", "PDFConfig"),
    ]
    
    for module_path, gen_class, config_class in import_paths:
        try:
            module = __import__(module_path, fromlist=[gen_class, config_class])
            _MedicalPDFGenerator = getattr(module, gen_class)
            _PDFConfig = getattr(module, config_class)
            logger.info(f"Loaded MedicalPDFGenerator from {module_path}")
            return _MedicalPDFGenerator, _PDFConfig
        except ImportError:
            continue
    
    raise ImportError(
        "MedicalPDFGenerator not found. Ensure medical_pdf_generator.py is in the Python path.\n"
        "Install dependencies: pip install reportlab pillow"
    )


# =============================================================================
# DATA STRUCTURES
# =============================================================================

@dataclass
class ResolvedFigure:
    """A figure that has been matched to an actual image."""
    figure_number: int
    placeholder_id: str
    figure_type: str
    topic: str
    image_path: Path
    caption: str
    confidence: float
    width_percent: int = 80


@dataclass 
class ExportConfig:
    """Configuration for export generation."""
    title: str = "Synthesis Output"
    author: str = "NeuroSynth"
    date: Optional[str] = None
    include_toc: bool = True
    include_cover: bool = True
    include_abstract: bool = True
    include_references: bool = True
    image_quality: str = "high"  # high, medium, low
    page_size: str = "letter"  # letter, a4
    primary_color: str = "#1a5f7a"
    accent_color: str = "#14b8a6"
    
    def __post_init__(self):
        if self.date is None:
            self.date = datetime.now().strftime("%B %d, %Y")


# =============================================================================
# FIGURE SUBSTITUTOR (for HTML/Markdown)
# =============================================================================

class FigureSubstitutor:
    """Resolves and substitutes figure placeholders with actual images."""
    
    PLACEHOLDER_PATTERN = re.compile(
        r'\[REQUEST_FIGURE:\s*type="([^"]+)"\s*topic="([^"]+)"\]',
        re.IGNORECASE
    )
    
    def __init__(self, image_base_path: Path):
        self.image_base_path = Path(image_base_path)
        self.figure_counter = 0
    
    def substitute_all(
        self,
        content: str,
        resolved_figures: List[Dict[str, Any]],
        output_format: str = "markdown"
    ) -> Tuple[str, List[ResolvedFigure]]:
        """Replace all placeholders with proper figure markup."""
        self.figure_counter = 0
        used_figures: List[ResolvedFigure] = []
        figures_by_keywords = self._build_keyword_index(resolved_figures)
        
        def replacer(match: re.Match) -> str:
            fig_type = match.group(1)
            topic = match.group(2)
            resolved = self._find_best_match(topic, fig_type, figures_by_keywords)
            
            self.figure_counter += 1
            if resolved:
                path = Path(resolved.get('path', resolved.get('storage_path', '')))
                figure = ResolvedFigure(
                    figure_number=self.figure_counter,
                    placeholder_id=f"fig_{self.figure_counter}",
                    figure_type=fig_type,
                    topic=topic,
                    image_path=path,
                    caption=resolved.get('caption', topic),
                    confidence=resolved.get('confidence', 0.8)
                )
                used_figures.append(figure)
                return self._format_figure(figure, output_format)
            else:
                return self._format_missing_figure(self.figure_counter, topic, output_format)
        
        substituted = self.PLACEHOLDER_PATTERN.sub(replacer, content)
        return substituted, used_figures
    
    def _build_keyword_index(self, figures: List[Dict]) -> Dict[str, List[Dict]]:
        """Index figures by their caption keywords for matching."""
        index: Dict[str, List[Dict]] = {}
        for fig in figures:
            caption = fig.get('caption', fig.get('topic', '')).lower()
            words = re.findall(r'\b\w{4,}\b', caption)
            for word in words:
                if word not in index:
                    index[word] = []
                index[word].append(fig)
        return index
    
    def _find_best_match(self, topic: str, fig_type: str, index: Dict) -> Optional[Dict]:
        """Find best matching figure for a placeholder."""
        topic_words = set(re.findall(r'\b\w{4,}\b', topic.lower()))
        
        candidates = []
        for word in topic_words:
            if word in index:
                for fig in index[word]:
                    fig_text = (fig.get('caption', '') + ' ' + fig.get('topic', '')).lower()
                    score = sum(1 for w in topic_words if w in fig_text)
                    candidates.append((score, fig))
        
        if candidates:
            candidates.sort(key=lambda x: x[0], reverse=True)
            return candidates[0][1]
        return None
    
    def _format_figure(self, figure: ResolvedFigure, fmt: str) -> str:
        """Format a resolved figure for the output format."""
        if fmt == "html":
            return f'''
<figure class="synthesis-figure" id="{figure.placeholder_id}">
    <img src="{{{{IMAGE:{figure.image_path}}}}}" alt="{figure.caption}">
    <figcaption><strong>Figure {figure.figure_number}:</strong> {figure.caption}</figcaption>
</figure>
'''
        else:  # markdown
            return f'''

![{figure.caption}]({figure.image_path})

**Figure {figure.figure_number}:** {figure.caption}

'''
    
    def _format_missing_figure(self, num: int, topic: str, fmt: str) -> str:
        """Format placeholder for missing figures."""
        if fmt == "html":
            return f'''
<figure class="synthesis-figure missing">
    <div class="placeholder">Image not available</div>
    <figcaption><strong>Figure {num}:</strong> {topic}</figcaption>
</figure>
'''
        return f"\n\n*[Figure {num}: {topic} — Image not available]*\n\n"


# =============================================================================
# HTML GENERATOR (for self-contained HTML output)
# =============================================================================

class HTMLGenerator:
    """Generates self-contained HTML with embedded base64 images."""
    
    CSS = '''
    :root {
        --primary: #1a5f7a;
        --accent: #14b8a6;
        --text: #1a1a2e;
        --muted: #666;
        --bg: #fff;
        --pearl-bg: #d4edda;
        --pearl-border: #28a745;
        --hazard-bg: #f8d7da;
        --hazard-border: #dc3545;
    }
    body {
        font-family: Georgia, 'Times New Roman', serif;
        max-width: 800px;
        margin: 40px auto;
        padding: 0 20px;
        line-height: 1.7;
        color: var(--text);
    }
    h1 { color: var(--primary); border-bottom: 2px solid var(--accent); padding-bottom: 10px; }
    h2 { color: var(--primary); margin-top: 30px; border-bottom: 1px solid #ddd; }
    .abstract { background: #f9f9f9; padding: 20px; border-left: 4px solid var(--accent); font-style: italic; }
    figure { margin: 30px 0; text-align: center; }
    figure img { max-width: 100%; border: 1px solid #ddd; border-radius: 4px; }
    figcaption { margin-top: 10px; font-size: 0.9em; color: var(--muted); }
    .pearl { background: var(--pearl-bg); border-left: 4px solid var(--pearl-border); padding: 15px; margin: 20px 0; }
    .pearl::before { content: "💡 PEARL: "; font-weight: bold; color: #155724; }
    .hazard { background: var(--hazard-bg); border-left: 4px solid var(--hazard-border); padding: 15px; margin: 20px 0; }
    .hazard::before { content: "⚠️ HAZARD: "; font-weight: bold; color: #721c24; }
    .reference { font-size: 0.9em; color: var(--muted); margin: 5px 0; padding-left: 30px; text-indent: -30px; }
    .meta { text-align: center; color: var(--muted); margin-bottom: 30px; }
    .word-count { text-align: right; font-size: 0.85em; color: var(--muted); font-style: italic; }
    '''
    
    def __init__(self, config: ExportConfig, image_base_path: Path):
        self.config = config
        self.image_base_path = Path(image_base_path)
    
    def generate(self, data: Dict[str, Any]) -> str:
        """Generate complete HTML document."""
        html_parts = [
            '<!DOCTYPE html>',
            '<html lang="en">',
            '<head>',
            '    <meta charset="UTF-8">',
            f'    <title>{data.get("title", "Synthesis")}</title>',
            f'    <style>{self.CSS}</style>',
            '</head>',
            '<body>',
            f'    <h1>{data.get("title", "Synthesis")}</h1>',
            f'    <div class="meta">Generated by {self.config.author} on {self.config.date}</div>',
        ]
        
        # Abstract
        if data.get('abstract'):
            html_parts.append(f'    <div class="abstract">{data["abstract"]}</div>')
        
        # Sections
        for section in data.get('sections', []):
            html_parts.append(f'    <h2>{section.get("title", "Section")}</h2>')
            content = self._process_content(section.get('content', ''))
            html_parts.append(f'    <div class="section-content">{content}</div>')
            word_count = section.get('word_count', 0)
            html_parts.append(f'    <div class="word-count">({word_count:,} words)</div>')
        
        # References
        if data.get('references'):
            html_parts.append('    <h2>References</h2>')
            for i, ref in enumerate(data['references']):
                source = ref.get('source', 'Unknown')
                html_parts.append(f'    <div class="reference">[{i+1}] {source}</div>')
        
        html_parts.extend(['</body>', '</html>'])
        
        html = '\n'.join(html_parts)
        
        # Embed images as base64
        html = self._embed_images(html)
        
        return html
    
    def _process_content(self, content: str) -> str:
        """Process content with pearl/hazard styling."""
        paragraphs = content.split('\n\n')
        processed = []
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Pearl
            if '[PEARL]' in para or '**PEARL:**' in para:
                text = re.sub(r'\[PEARL\]\s*|\*\*PEARL:\*\*\s*', '', para)
                processed.append(f'<div class="pearl">{text}</div>')
            # Hazard
            elif '[HAZARD]' in para or '**HAZARD:**' in para:
                text = re.sub(r'\[HAZARD\]\s*|\*\*HAZARD:\*\*\s*', '', para)
                processed.append(f'<div class="hazard">{text}</div>')
            # Figure placeholder (already processed by substitutor)
            elif '<figure' in para:
                processed.append(para)
            else:
                processed.append(f'<p>{para}</p>')
        
        return '\n'.join(processed)
    
    def _embed_images(self, html: str) -> str:
        """Replace image placeholders with base64 data URIs."""
        pattern = re.compile(r'\{\{IMAGE:([^}]+)\}\}')
        
        def replacer(match):
            path_str = match.group(1)
            path = Path(path_str)
            if not path.is_absolute():
                path = self.image_base_path / path
            
            if path.exists():
                try:
                    with open(path, 'rb') as f:
                        data = base64.b64encode(f.read()).decode()
                    suffix = path.suffix.lower()
                    mime = {'jpg': 'jpeg', 'jpeg': 'jpeg', 'png': 'png', 'gif': 'gif'}.get(suffix[1:], 'png')
                    return f'data:image/{mime};base64,{data}'
                except Exception as e:
                    logger.warning(f"Failed to embed image {path}: {e}")
            
            # Placeholder SVG
            return 'data:image/svg+xml;base64,' + base64.b64encode(
                b'<svg xmlns="http://www.w3.org/2000/svg" width="400" height="200">'
                b'<rect fill="#f0f0f0" width="400" height="200"/>'
                b'<text x="200" y="100" text-anchor="middle" fill="#999">Image not found</text>'
                b'</svg>'
            ).decode()
        
        return pattern.sub(replacer, html)


# =============================================================================
# DOCX GENERATOR
# =============================================================================

class DOCXGenerator:
    """Generates Word documents with embedded images."""
    
    def __init__(self, config: ExportConfig, image_base_path: Path):
        self.config = config
        self.image_base_path = Path(image_base_path)
    
    def generate(self, data: Dict[str, Any], output_path: Path) -> None:
        """Generate DOCX file."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx required for DOCX generation. Install: pip install python-docx")
        
        doc = Document()
        
        # Title
        title_para = doc.add_heading(data.get('title', 'Synthesis'), level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Meta
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Generated by {self.config.author}\n").italic = True
        meta.add_run(self.config.date).italic = True
        
        doc.add_page_break()
        
        # Abstract
        if data.get('abstract'):
            doc.add_heading("Abstract", level=1)
            doc.add_paragraph(data['abstract']).italic = True
        
        # Sections
        figure_counter = 0
        resolved_figures = {
            f.get('topic', '').lower(): f 
            for f in data.get('resolved_figures', [])
        }
        
        for section in data.get('sections', []):
            doc.add_heading(section.get('title', 'Section'), level=1)
            
            content = section.get('content', '')
            paragraphs = content.split('\n\n')
            
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                
                # Figure placeholder
                fig_match = re.search(r'\[REQUEST_FIGURE:[^\]]+topic="([^"]+)"', para)
                if fig_match:
                    topic = fig_match.group(1)
                    figure_counter += 1
                    
                    # Find matching image
                    fig_data = None
                    for key, fig in resolved_figures.items():
                        if any(word in key for word in topic.lower().split()[:2]):
                            fig_data = fig
                            break
                    
                    if fig_data:
                        img_path = self.image_base_path / fig_data.get('path', fig_data.get('storage_path', ''))
                        if img_path.exists():
                            try:
                                doc.add_picture(str(img_path), width=Inches(5))
                                caption = doc.add_paragraph(f"Figure {figure_counter}: {fig_data.get('caption', topic)}")
                                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            except Exception as e:
                                doc.add_paragraph(f"[Figure {figure_counter}: {topic} — Could not load image]")
                    continue
                
                # Pearl
                if '[PEARL]' in para or 'PEARL:' in para:
                    text = re.sub(r'\[PEARL\]\s*|\*\*PEARL:\*\*\s*', '', para)
                    p = doc.add_paragraph()
                    p.add_run("💡 PEARL: ").bold = True
                    p.add_run(text)
                    continue
                
                # Hazard
                if '[HAZARD]' in para or 'HAZARD:' in para:
                    text = re.sub(r'\[HAZARD\]\s*|\*\*HAZARD:\*\*\s*', '', para)
                    p = doc.add_paragraph()
                    p.add_run("⚠️ HAZARD: ").bold = True
                    p.add_run(text)
                    continue
                
                # Regular paragraph
                doc.add_paragraph(para)
        
        # References
        if data.get('references'):
            doc.add_heading("References", level=1)
            for i, ref in enumerate(data['references']):
                doc.add_paragraph(f"[{i+1}] {ref.get('source', 'Unknown')}")
        
        doc.save(str(output_path))
        logger.info(f"DOCX generated: {output_path}")


# =============================================================================
# MAIN EXPORTER CLASS
# =============================================================================

class SynthesisExporter:
    """
    Unified export interface for synthesis results.
    
    Routes to optimal generator for each format:
    - PDF: MedicalPDFGenerator (ReportLab) - best quality, direct image embedding
    - HTML: HTMLGenerator with base64 embedding
    - DOCX: DOCXGenerator with python-docx
    - Markdown: FigureSubstitutor for link resolution
    
    Usage:
        exporter = SynthesisExporter(image_base_path=Path("data/images"))
        
        # Best quality PDF
        exporter.to_pdf(synthesis_result, Path("output.pdf"))
        
        # Self-contained HTML
        html = exporter.to_html(synthesis_result)
        
        # Editable Word document
        exporter.to_docx(synthesis_result, Path("output.docx"))
        
        # Markdown with image links
        md = exporter.to_markdown(synthesis_result)
    """
    
    def __init__(
        self,
        image_base_path: Path = Path("data/images"),
        config: Optional[ExportConfig] = None
    ):
        self.image_base_path = Path(image_base_path)
        self.config = config or ExportConfig()
        self.substitutor = FigureSubstitutor(self.image_base_path)
    
    def to_pdf(self, result: Any, output_path: Path) -> None:
        """
        Export to high-quality PDF using MedicalPDFGenerator (ReportLab).
        
        This is the optimal method for publication-ready output:
        - Direct image embedding (no base64 conversion)
        - Lossless image quality
        - Native clinical callout rendering (pearls/hazards)
        - Professional typography
        - Automatic figure numbering and placement
        
        Args:
            result: SynthesisResult object or dict
            output_path: Output PDF file path
        """
        # Get the ReportLab generator
        MedicalPDFGenerator, PDFConfig = _get_pdf_generator()
        
        data = self._normalize_result(result)
        
        # Configure PDF generator
        pdf_config = PDFConfig(
            author=self.config.author,
            include_toc=self.config.include_toc,
            include_cover=self.config.include_cover,
        )
        
        # Set colors if PDFConfig supports it
        try:
            from reportlab.lib.colors import HexColor
            if hasattr(pdf_config, 'primary_color') and self.config.primary_color:
                pdf_config.primary_color = HexColor(self.config.primary_color)
            if hasattr(pdf_config, 'accent_color') and self.config.accent_color:
                pdf_config.accent_color = HexColor(self.config.accent_color)
        except ImportError:
            pass  # Colors will use defaults
        
        # Create generator and build PDF
        generator = MedicalPDFGenerator(
            config=pdf_config,
            image_base_path=self.image_base_path
        )
        
        generator.generate(data, output_path)
        logger.info(f"PDF generated via MedicalPDFGenerator: {output_path}")
    
    def to_html(self, result: Any, embed_images: bool = True) -> str:
        """
        Export to self-contained HTML.
        
        Args:
            result: SynthesisResult object or dict
            embed_images: If True, embed images as base64 (default)
            
        Returns:
            HTML string
        """
        data = self._normalize_result(result)
        
        # Process sections to substitute figure placeholders
        processed_sections = []
        for section in data.get('sections', []):
            content = section.get('content', '')
            content, _ = self.substitutor.substitute_all(
                content,
                data.get('resolved_figures', []),
                output_format="html"
            )
            processed_sections.append({
                **section,
                'content': content
            })
        
        data['sections'] = processed_sections
        
        # Generate HTML
        generator = HTMLGenerator(self.config, self.image_base_path)
        return generator.generate(data)
    
    def to_docx(self, result: Any, output_path: Path) -> None:
        """
        Export to Microsoft Word document.
        
        Args:
            result: SynthesisResult object or dict
            output_path: Output DOCX file path
        """
        data = self._normalize_result(result)
        generator = DOCXGenerator(self.config, self.image_base_path)
        generator.generate(data, output_path)
    
    def to_markdown(self, result: Any) -> str:
        """
        Export to Markdown with resolved figure links.
        
        Args:
            result: SynthesisResult object or dict
            
        Returns:
            Markdown string
        """
        data = self._normalize_result(result)
        
        md_parts = [
            f"# {data.get('title', 'Synthesis')}\n",
            f"*Generated by {self.config.author} on {self.config.date}*\n",
            "\n---\n\n",
        ]
        
        # Abstract
        if data.get('abstract'):
            md_parts.append(f"## Abstract\n\n{data['abstract']}\n\n---\n\n")
        
        # Sections
        for section in data.get('sections', []):
            title = section.get('title', 'Section')
            content = section.get('content', '')
            word_count = section.get('word_count', len(content.split()))
            
            # Substitute figure placeholders
            content, _ = self.substitutor.substitute_all(
                content,
                data.get('resolved_figures', []),
                output_format="markdown"
            )
            
            md_parts.append(f"## {title}\n\n")
            md_parts.append(content)
            md_parts.append(f"\n*({word_count:,} words)*\n\n---\n\n")
        
        # References
        if data.get('references'):
            md_parts.append("## References\n\n")
            for i, ref in enumerate(data['references']):
                source = ref.get('source', 'Unknown')
                md_parts.append(f"{i+1}. {source}\n")
        
        return ''.join(md_parts)
    
    def _normalize_result(self, result: Any) -> Dict[str, Any]:
        """Convert result object to dict format."""
        if isinstance(result, dict):
            return result
        
        # Handle SynthesisResult-like objects
        return {
            'title': getattr(result, 'title', 'Untitled'),
            'abstract': getattr(result, 'abstract', ''),
            'sections': [
                s.to_dict() if hasattr(s, 'to_dict') else s
                for s in getattr(result, 'sections', [])
            ],
            'references': getattr(result, 'references', []),
            'resolved_figures': getattr(result, 'resolved_figures', []),
            'metadata': getattr(result, 'metadata', {})
        }


# =============================================================================
# CONVENIENCE FUNCTIONS
# =============================================================================

def export_to_pdf(
    synthesis_data: Dict[str, Any],
    output_path: Path,
    image_base_path: Path = Path("data/images"),
    author: str = "NeuroSynth"
) -> None:
    """Quick function to export synthesis to PDF."""
    config = ExportConfig(author=author)
    exporter = SynthesisExporter(image_base_path=image_base_path, config=config)
    exporter.to_pdf(synthesis_data, output_path)


def export_to_html(
    synthesis_data: Dict[str, Any],
    output_path: Path,
    image_base_path: Path = Path("data/images")
) -> None:
    """Quick function to export synthesis to HTML file."""
    exporter = SynthesisExporter(image_base_path=image_base_path)
    html = exporter.to_html(synthesis_data)
    with open(output_path, 'w') as f:
        f.write(html)


def export_to_markdown(
    synthesis_data: Dict[str, Any],
    output_path: Path,
    image_base_path: Path = Path("data/images")
) -> None:
    """Quick function to export synthesis to Markdown file."""
    exporter = SynthesisExporter(image_base_path=image_base_path)
    md = exporter.to_markdown(synthesis_data)
    with open(output_path, 'w') as f:
        f.write(md)


# =============================================================================
# CLI
# =============================================================================

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Export synthesis to various formats")
    parser.add_argument("input", type=Path, help="Input JSON file")
    parser.add_argument("output", type=Path, help="Output file (.pdf, .html, .docx, .md)")
    parser.add_argument("--images", type=Path, default=Path("data/images"), help="Image base path")
    parser.add_argument("--author", type=str, default="NeuroSynth", help="Author name")
    
    args = parser.parse_args()
    
    # Load data
    with open(args.input, 'r') as f:
        data = json.load(f)
    
    # Configure
    config = ExportConfig(
        title=data.get('title', 'Synthesis'),
        author=args.author
    )
    
    exporter = SynthesisExporter(image_base_path=args.images, config=config)
    
    # Export based on suffix
    suffix = args.output.suffix.lower()
    
    if suffix == '.pdf':
        exporter.to_pdf(data, args.output)
        print(f"✅ PDF: {args.output}")
    elif suffix == '.html':
        html = exporter.to_html(data)
        with open(args.output, 'w') as f:
            f.write(html)
        print(f"✅ HTML: {args.output}")
    elif suffix == '.docx':
        exporter.to_docx(data, args.output)
        print(f"✅ DOCX: {args.output}")
    elif suffix in ('.md', '.markdown'):
        md = exporter.to_markdown(data)
        with open(args.output, 'w') as f:
            f.write(md)
        print(f"✅ Markdown: {args.output}")
    else:
        print(f"❌ Unknown format: {suffix}")
        print("Supported: .pdf, .html, .docx, .md")
